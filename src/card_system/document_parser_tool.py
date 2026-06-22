"""File parsing tool used by CardReviewAgent upload analysis."""

from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any


class DocumentParseError(ValueError):
    """Raised when an uploaded learning document cannot be parsed."""


class DocumentParserTool:
    """Parse PDF, DOCX, TXT and Markdown uploads into plain text."""

    supported_extensions = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    def parse(self, filename: str, content: bytes) -> dict[str, Any]:
        suffix = Path(filename).suffix.lower()
        if suffix not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise DocumentParseError(f"不支持的文件类型：{suffix or '<none>'}。支持：{supported}")
        if not content:
            raise DocumentParseError("上传文件为空，无法解析学习资料。")

        if suffix == ".pdf":
            text = self._parse_pdf(content)
            file_type = "pdf"
        elif suffix == ".docx":
            text = self._parse_docx(content)
            file_type = "docx"
        elif suffix in {".md", ".markdown"}:
            text = self._decode_text(content)
            file_type = "markdown"
        else:
            text = self._decode_text(content)
            file_type = "txt"

        text = self._normalize_text(text)
        if not text:
            raise DocumentParseError("文件中没有提取到可用于学习分析的文本。")

        return {
            "filename": filename,
            "file_type": file_type,
            "text": text,
            "char_count": len(text),
        }

    def _parse_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise DocumentParseError("PDF 解析需要 pypdf，请先安装项目依赖。") from e

        try:
            reader = PdfReader(BytesIO(content))
            return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
        except Exception as e:
            raise DocumentParseError("PDF 文件解析失败，请确认文件可读取且未加密。") from e

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document

            document = Document(BytesIO(content))
            parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            return self._parse_docx_with_docx2txt(content)
        except Exception:
            fallback = self._parse_docx_with_docx2txt(content)
            if fallback:
                return fallback
            raise DocumentParseError("DOCX 文件解析失败，请确认文件格式正确。")

    def _parse_docx_with_docx2txt(self, content: bytes) -> str:
        try:
            import docx2txt
        except ImportError as e:
            raise DocumentParseError("DOCX 解析需要 python-docx 或 docx2txt，请先安装项目依赖。") from e

        temp_path: Path | None = None
        try:
            with NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(content)
                temp_path = Path(temp_file.name)
            return docx2txt.process(str(temp_path)) or ""
        except Exception as e:
            raise DocumentParseError("DOCX 文件解析失败，请确认文件格式正确。") from e
        finally:
            if temp_path:
                temp_path.unlink(missing_ok=True)

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("文本文件编码无法识别。")

    def _normalize_text(self, text: str) -> str:
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line).strip()
